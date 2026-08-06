from __future__ import annotations
import hashlib
# --- Set-based idempotency helper ---
def calculate_claim_set_hash(claim_id, db):
    """Fetch all content_hash for claim, sort, join, and return SHA-256 hash."""
    hashes = [d.content_hash for d in db.query(Document).filter(Document.claim_id == claim_id).all() if d.content_hash]
    hashes.sort()
    joined = ",".join(hashes)
    return hashlib.sha256(joined.encode("utf-8")).hexdigest()

import hashlib
import logging
import os
import re
import sys
import sys
import uuid
from datetime import datetime
from pathlib import Path, PurePosixPath
from typing import Any

import aiofiles
from celery import chord, group, chain
from fastapi import APIRouter, Body, Depends, FastAPI, File, Form, HTTPException, Query, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from fastapi.exceptions import RequestValidationError
from starlette.exceptions import HTTPException as StarletteHTTPException
from services.shared_tasks import (
    coding_task,
    intake_task,
    ocr_task,
    parser_task,
    risk_task,
    validator_task,
    finalize_claim_task,
    run_pipeline_inline,
)
from libs.shared.celery_app import celery_app
from pydantic import BaseModel, EmailStr, Field
from sqlalchemy import text
from sqlalchemy.orm import Session, selectinload

from .config import settings
from .db import SessionLocal, check_db_health, engine, force_master_session
from .models import Claim, Document, DocValidation
from libs.auth.passwords import hash_password, password_matches, verify_password
from libs.shared.models import ParseJob, ParsedField, WorkflowState
from libs.shared.workflow_state import get_latest_workflow_state, upsert_workflow_state
from .schemas import ClaimListOut, ClaimOut


try:
    from libs.utils.audit import AuditLogger
except Exception:
    AuditLogger = None  # type: ignore

def _audit(db, action: str, claim_id=None, metadata=None):
    try:
        if AuditLogger:
            with SessionLocal() as audit_db:
                AuditLogger(audit_db, "ingress").log(action, claim_id=claim_id, metadata=metadata)
    except Exception:
        logger.debug("Audit log failed for %s", action, exc_info=True)

# ------------------------------------------------------------------ logging
logging.basicConfig(
    level=settings.log_level.upper(),
    format="%(asctime)s  %(levelname)-8s  %(name)s  %(message)s",
)
logger = logging.getLogger("ingress")

# log4net-style on-disk audit log for claim uploads
# Writes to <repo_root>/logs/claim_uploads.txt (override via CLAIMGPT_LOG_DIR).
try:
    sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "..", ".."))
    from libs.observability.file_logger import get_file_logger
    upload_log = get_file_logger("ingress.upload", "claim_uploads.txt")
except Exception:  # pragma: no cover - logging must never break the service
    logger.exception("Failed to initialise claim upload file logger; falling back to standard logger")
    upload_log = logger

RAW_STORAGE = Path(settings.storage_root).resolve()
RAW_STORAGE.mkdir(parents=True, exist_ok=True)

app = FastAPI(title="ClaimGPT Ingress Service")

# Global exception handler to ensure all errors return JSON
@app.exception_handler(Exception)
async def global_exception_handler(request, exc):
    logger.exception("Unhandled exception in ingress service")
    return JSONResponse(
        status_code=500,
        content={"detail": f"Internal server error: {str(exc)}"},
    )

@app.exception_handler(StarletteHTTPException)
async def http_exception_handler(request, exc):
    return JSONResponse(
        status_code=exc.status_code,
        content={"detail": exc.detail},
    )

@app.exception_handler(RequestValidationError)
async def validation_exception_handler(request, exc):
    return JSONResponse(
        status_code=422,
        content={"detail": str(exc)},
    )

# ------------------------------------------------------------------ CORS
app.add_middleware(
    CORSMiddleware,
    allow_origin_regex="https://.*|http://localhost:.*|http://127.0.0.1:.*",
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ------------------------------------------------------------------ observability
try:
    import sys
    sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "..", ".."))
    from libs.observability.metrics import PrometheusMiddleware, init_metrics, metrics_endpoint
    from libs.observability.tracing import init_tracing, instrument_fastapi
    init_tracing("ingress")
    init_metrics("ingress")
    instrument_fastapi(app)
    app.add_middleware(PrometheusMiddleware)
    _metrics_handler = metrics_endpoint()
    if _metrics_handler:
        app.get("/metrics")(_metrics_handler)
except Exception:
    logger.debug("Observability libs not available — skipping")


# ------------------------------------------------------------------ lifecycle
@app.on_event("shutdown")
def _shutdown():
    engine.dispose()
    logger.info("DB engine disposed")


# ------------------------------------------------------------------ deps
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


def _parse_uuid(value: str) -> uuid.UUID:
    try:
        return uuid.UUID(value)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid UUID")


def _safe_filename(raw: str | None) -> str:
    """Strip directory components to prevent path-traversal via filename."""
    if not raw:
        return "upload.bin"
    return PurePosixPath(raw).name or "upload.bin"


# Map every file extension we accept to one canonical Content-Type so we can
# normalise uploads coming from clients that send non-standard MIMEs (e.g.
# Windows reporting ``image/jpg`` for .jpg, or browsers/curl falling back to
# ``application/octet-stream``).  Keep this in lock-step with the OCR engine's
# SUPPORTED_EXTENSIONS — anything OCR can read should be uploadable.
_EXTENSION_TO_CONTENT_TYPE: dict[str, str] = {
    ".pdf": "application/pdf",
    # Images
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".jpe": "image/jpeg",
    ".jfif": "image/jpeg",
    ".png": "image/png",
    ".tiff": "image/tiff",
    ".tif": "image/tiff",
    ".bmp": "image/bmp",
    ".webp": "image/webp",
    ".gif": "image/gif",
    ".heic": "image/heic",
    ".heif": "image/heif",
    # Office
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".doc": "application/msword",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".xls": "application/vnd.ms-excel",
    ".pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
    ".ppt": "application/vnd.ms-powerpoint",
    # OpenDocument
    ".odt": "application/vnd.oasis.opendocument.text",
    ".ods": "application/vnd.oasis.opendocument.spreadsheet",
    ".odp": "application/vnd.oasis.opendocument.presentation",
    # Misc
    ".rtf": "application/rtf",
    ".txt": "text/plain",
    ".csv": "text/csv",
    ".json": "application/json",
    ".xml": "application/xml",
    ".html": "text/html",
    ".htm": "text/html",
}

# Common non-standard / aliased MIME types we should accept silently.
_CONTENT_TYPE_ALIASES: dict[str, str] = {
    "image/jpg": "image/jpeg",       # non-standard but seen in the wild (Windows)
    "image/pjpeg": "image/jpeg",     # progressive JPEG (legacy IE)
    "image/x-png": "image/png",      # legacy
    "image/x-citrix-jpeg": "image/jpeg",
    "image/x-citrix-png": "image/png",
    "text/xml": "application/xml",
}


def _resolve_content_type(file: UploadFile) -> tuple[str, bool]:
    """Decide the effective Content-Type for an upload.

    Returns ``(content_type, is_supported)``.  Falls back to the file extension
    when the client sends nothing useful (``application/octet-stream`` or an
    empty header).  This is the single source of truth for upload validation
    so `.jpg` files always pass even when browsers report `image/jpg`.
    """
    raw_ct = (file.content_type or "").lower().strip()
    suffix = Path(file.filename or "").suffix.lower()

    # 1) Direct match against allowed list.
    if raw_ct in settings.allowed_content_types:
        return raw_ct, True

    # 2) Try alias normalisation.
    if raw_ct in _CONTENT_TYPE_ALIASES:
        canonical = _CONTENT_TYPE_ALIASES[raw_ct]
        if canonical in settings.allowed_content_types:
            return canonical, True

    # 3) Browsers / curl often send application/octet-stream or nothing for
    #    unknown extensions — trust the file extension as long as we know it.
    if suffix in _EXTENSION_TO_CONTENT_TYPE:
        canonical = _EXTENSION_TO_CONTENT_TYPE[suffix]
        if canonical in settings.allowed_content_types:
            return canonical, True

    return raw_ct or "application/octet-stream", False


def _compute_upload_sha256(file_data: list[tuple[UploadFile, bytes, str]]) -> str:
    hasher = hashlib.sha256()
    for _, content, safe_name in file_data:
        hasher.update(safe_name.encode("utf-8", errors="ignore"))
        hasher.update(b"\x00")
        hasher.update(content)
        hasher.update(b"\x00")
    return hasher.hexdigest()


def _build_claim_response(db: Session, claim_id: uuid.UUID, extra: dict[str, Any] | None = None) -> dict[str, Any]:
    claim = (
        db.query(Claim)
        .options(selectinload(Claim.documents))
        .filter(Claim.id == claim_id)
        .first()
    )
    if not claim:
        raise HTTPException(status_code=404, detail="Claim not found")
    payload = ClaimOut.model_validate(claim).model_dump(mode="json")
    if extra:
        payload.update(extra)
    return payload


def _build_report_url(claim_id: uuid.UUID) -> str:
    return f"/claims/{claim_id}"


def _find_completed_claim_by_upload_hash(db: Session, upload_sha256: str) -> Claim | None:
    row = db.execute(
        text(
            """
            SELECT c.id
            FROM claims c
            JOIN audit_logs a ON a.claim_id = c.id
            WHERE a.action = 'CLAIM_CREATED'
              AND a.metadata->>'upload_sha256' = :upload_sha256
              AND c.status = 'COMPLETED'
            ORDER BY c.created_at DESC
            LIMIT 1
            """
        ),
        {"upload_sha256": upload_sha256},
    ).first()
    if not row:
        return None
    return db.query(Claim).filter(Claim.id == row[0]).first()


def _celery_worker_available(timeout: float = 0.6) -> bool:
    """Best-effort check that at least one Celery worker is online and ready.

    A short ping (<1s) is issued via the Celery control bus. If the broker is
    unreachable or no worker replies within the timeout, returns ``False`` —
    callers can then fall back to inline execution so uploads never get stuck.
    """
    try:
        replies = celery_app.control.ping(timeout=timeout) or []
        return bool(replies)
    except Exception:
        return False


def _should_run_inline() -> bool:
    """Decide between the Celery chain and in-process inline execution.

    Resolution order:
      * ``CLAIMGPT_INLINE_PIPELINE=1`` / ``true`` / ``yes``  -> always inline
      * ``CLAIMGPT_INLINE_PIPELINE=0`` / ``false`` / ``no``  -> never inline (require worker)
      * ``CLAIMGPT_INLINE_PIPELINE`` unset or ``auto``       -> inline only if no worker is reachable
    """
    raw = (os.getenv("CLAIMGPT_INLINE_PIPELINE") or "auto").strip().lower()
    if raw in {"1", "true", "yes", "on", "inline"}:
        return True
    if raw in {"0", "false", "no", "off", "celery"}:
        return False
    # auto: inline only when no worker is online
    return not _celery_worker_available()


def _enqueue_pipeline(
    file_metadata: list[dict[str, str]] | str,
    policy_id: str | None = None,
    patient_id: str | None = None,
) -> str:
    """Enqueue the full pipeline starting with intake task, or trigger OCR on an existing claim.
    
    Args:
        file_metadata: List of dicts with keys: path, safe_name, content_hash, effective_ct
                      OR a string claim_id to retrigger pipeline for an existing claim.
        policy_id: Optional policy ID
        patient_id: Optional patient ID
    
    Returns:
        Task ID as string, or "inline:{claim_id}" for inline execution
    """
    if isinstance(file_metadata, str):
        claim_id_str = file_metadata
        
        # Synchronously reset workflow state in DB so polling API immediately returns 5% progress
        try:
            from services.ocr.app.db import SessionLocal as OcrSessionLocal
            with OcrSessionLocal() as db_session:
                upsert_workflow_state(db_session, uuid.UUID(claim_id_str), "STARTING", status="RUNNING")
                db_session.commit()
        except Exception:
            logger.exception("Failed to reset workflow state for claim %s", claim_id_str)

        if _should_run_inline():
            import threading
            logger.warning(
                "Celery worker not detected (or inline mode forced) — running pipeline inline for existing claim %s",
                claim_id_str,
            )

            def _runner() -> None:
                try:
                    run_pipeline_inline(claim_id_str)
                except Exception:
                    logger.exception("Inline pipeline crashed")

            thread = threading.Thread(
                target=_runner,
                name="inline-pipeline",
                daemon=True,
            )
            thread.start()
            return "inline:queued"

        workflow_chain = chain(
            ocr_task.s(claim_id_str),                               # Step 2: OCR (intake bypassed)
            parser_task.s(),                                        # Step 3: Parser
            coding_task.s(),                                        # Step 4: Coding
            risk_task.s(),                                          # Step 5: Risk
            validator_task.s(),                                     # Step 6: Validator
            finalize_claim_task.s(),                                # Step 7: Finalize Callback
        )
        result = workflow_chain.apply_async()
        return str(result.id)

    if _should_run_inline():
        # For inline execution, intake task needs to create the claim first
        import threading
        logger.warning(
            "Celery worker not detected (or inline mode forced) — running pipeline inline",
        )

        def _runner() -> None:
            try:
                # For inline, we need to do intake synchronously first
                from services.ocr.app.db import SessionLocal as OcrSessionLocal
                import hashlib
                db = OcrSessionLocal()
                try:
                    # Create claim
                    claim = Claim(
                        policy_id=policy_id,
                        patient_id=patient_id,
                        status="UPLOADED",
                        source="PATIENT",
                    )
                    db.add(claim)
                    db.flush()
                    claim_id = claim.id
                    
                    # Create documents
                    for metadata in file_metadata:
                        doc = Document(
                            claim_id=claim_id,
                            file_name=metadata["safe_name"],
                            file_type=metadata["effective_ct"],
                            minio_path=metadata["path"],
                            content_hash=metadata["content_hash"],
                        )
                        db.add(doc)
                    
                    db.commit()
                    
                    # Calculate set_hash
                    hashes = [d.content_hash for d in db.query(Document).filter(Document.claim_id == claim_id).all() if d.content_hash]
                    hashes.sort()
                    set_hash = hashlib.sha256(",".join(hashes).encode("utf-8")).hexdigest()
                    
                    # Create ParseJob
                    from libs.shared.models import ParseJob as PJ
                    parse_job = PJ(claim_id=claim_id, status="PENDING", set_hash=set_hash)
                    db.add(parse_job)
                    db.commit()
                    
                    # Update workflow state
                    upsert_workflow_state(db, claim_id, "STARTING", status="RUNNING")
                    db.commit()
                    
                    claim_id_str = str(claim_id)
                finally:
                    db.close()
                
                # Now run the inline pipeline
                run_pipeline_inline(claim_id_str)
            except Exception:
                logger.exception("Inline pipeline crashed")

        thread = threading.Thread(
            target=_runner,
            name="inline-pipeline",
            daemon=True,
        )
        thread.start()
        return "inline:queued"

    workflow_chain = chain(
        intake_task.s(file_metadata, policy_id, patient_id),  # Step 1: Intake (DB operations)
        ocr_task.s(),                                           # Step 2: OCR
        parser_task.s(),                                        # Step 3: Parser
        coding_task.s(),                                        # Step 4: Coding
        risk_task.s(),                                          # Step 5: Risk
        validator_task.s(),                                     # Step 6: Validator
        finalize_claim_task.s(),                                # Step 7: Finalize Callback
    )
    result = workflow_chain.apply_async()
    return str(result.id)


def _get_step_index(current_step: str | None, status: str | None) -> int:
    if current_step in ['OCR_STARTED', 'OCR_FINISHED']:
        return 1
    elif current_step in ['PARSING_STARTED', 'PARSING_FINISHED']:
        return 2
    elif current_step in ['CODING_STARTED', 'CODING_FINISHED', 'RISK_STARTED', 'RISK_FINISHED', 'VALIDATION_STARTED', 'VALIDATION_FINISHED']:
        return 3
    elif current_step in ['FINALIZE_STARTED', 'FINALIZE_FINISHED']:
        return 4
    elif status == 'FINISHED':
        return 5
    else:
        return 0


_PATIENT_NAME_PATTERNS = [
    re.compile(r"(?im)(?:^|\n)\s*(?:patient\s*name|name\s*of\s*patient)\s*[:\-]\s*([^\n\r|]+)"),
]

_DOB_PATTERNS = [
    re.compile(r"(?im)(?:^|\n)\s*(?:date\s*of\s*birth|dob|d\.o\.b)\s*[:\-]\s*([^\n\r|]+)"),
]

_MONTHS = {
    "jan": 1, "january": 1,
    "feb": 2, "february": 2,
    "mar": 3, "march": 3,
    "apr": 4, "april": 4,
    "may": 5,
    "jun": 6, "june": 6,
    "jul": 7, "july": 7,
    "aug": 8, "august": 8,
    "sep": 9, "sept": 9, "september": 9,
    "oct": 10, "october": 10,
    "nov": 11, "november": 11,
    "dec": 12, "december": 12,
}


def _canonical_name(value: str | None) -> str:
    if not value:
        return ""
    return re.sub(r"\s+", " ", value).strip().lower()


def _normalize_dob(value: str | None) -> str:
    if not value:
        return ""
    raw = re.sub(r"\s+", " ", value).strip().replace(",", "")
    m_num = re.fullmatch(r"(\d{1,2})[\-/.](\d{1,2})[\-/.](\d{2,4})", raw)
    if m_num:
        day, month, year = int(m_num.group(1)), int(m_num.group(2)), int(m_num.group(3))
        if year < 100:
            year += 2000 if year < 50 else 1900
        return f"{year:04d}-{month:02d}-{day:02d}"

    m_mon = re.fullmatch(r"(\d{1,2})[\-/. ]([A-Za-z]{3,9})[\-/. ](\d{2,4})", raw)
    if m_mon:
        day, month_token, year = int(m_mon.group(1)), m_mon.group(2).lower(), int(m_mon.group(3))
        month = _MONTHS.get(month_token)
        if month:
            if year < 100:
                year += 2000 if year < 50 else 1900
            return f"{year:04d}-{month:02d}-{day:02d}"

    m_alt = re.fullmatch(r"([A-Za-z]{3,9})\s+(\d{1,2})\s+(\d{2,4})", raw)
    if m_alt:
        month_token, day, year = m_alt.group(1).lower(), int(m_alt.group(2)), int(m_alt.group(3))
        month = _MONTHS.get(month_token)
        if month:
            if year < 100:
                year += 2000 if year < 50 else 1900
            return f"{year:04d}-{month:02d}-{day:02d}"

    return raw.lower()


def _extract_text_for_identity(file_path: Path, file_type: str | None) -> str:
    file_type = (file_type or "").lower()
    suffix = file_path.suffix.lower()

    if file_type == "application/pdf" or suffix == ".pdf":
        try:
            import pdfplumber

            parts: list[str] = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages[:5]:
                    t = page.extract_text() or ""
                    if t.strip():
                        parts.append(t)
            return "\n".join(parts)
        except Exception:
            return ""

    if suffix == ".docx":
        try:
            import docx

            d = docx.Document(str(file_path))
            return "\n".join(p.text for p in d.paragraphs if p.text)
        except Exception:
            return ""

    if suffix in {".xlsx", ".xlsm"}:
        try:
            import openpyxl

            wb = openpyxl.load_workbook(str(file_path), read_only=True, data_only=True)
            lines: list[str] = []
            for ws in wb.worksheets[:3]:
                for row in ws.iter_rows(min_row=1, max_row=60, values_only=True):
                    vals = [str(c).strip() for c in row if c is not None and str(c).strip()]
                    if vals:
                        lines.append(" | ".join(vals))
            return "\n".join(lines)
        except Exception:
            return ""

    if suffix in {".txt", ".csv", ".json", ".xml", ".html", ".htm"}:
        try:
            return file_path.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            return ""

    return ""


def _extract_identity_from_text(text: str) -> tuple[str | None, str | None]:
    if not text:
        return None, None

    patient_name: str | None = None
    dob: str | None = None

    for pat in _PATIENT_NAME_PATTERNS:
        m = pat.search(text)
        if m:
            patient_name = m.group(1).strip()
            break

    for pat in _DOB_PATTERNS:
        m = pat.search(text)
        if m:
            dob = m.group(1).strip()
            break

    if patient_name:
        patient_name = re.sub(r"\s+", " ", patient_name).strip()
    if dob:
        dob = re.sub(r"\s+", " ", dob).strip()
    return patient_name, dob


def _existing_identity_anchor(db: Session, claim_id: uuid.UUID) -> tuple[str | None, str | None]:
    rows = (
        db.query(DocValidation)
        .filter(
            DocValidation.claim_id == claim_id,
            DocValidation.doc_type == "IDENTITY_GATE",
            DocValidation.status == "VALID",
        )
        .order_by(DocValidation.created_at.asc())
        .all()
    )
    if rows:
        locked = []
        for row in rows:
            md = row.validation_metadata or {}
            if md.get("anchor_locked"):
                locked.append(row)
        picked = locked[0] if locked else rows[0]
        md = picked.validation_metadata or {}
        return picked.patient_name, md.get("identity_dob")

    # Fallback to other DocValidation rows (e.g. from the first batch OCR/validation)
    other_val = (
        db.query(DocValidation)
        .filter(
            DocValidation.claim_id == claim_id,
            DocValidation.status == "VALID",
            DocValidation.patient_name.isnot(None),
        )
        .order_by(DocValidation.created_at.asc())
        .first()
    )
    if other_val:
        md = other_val.validation_metadata or {}
        return other_val.patient_name, md.get("identity_dob")

    # Fallback to ParsedField (populated by LLM parser for first batch)
    pf_name = (
        db.query(ParsedField.field_value)
        .filter(
            ParsedField.claim_id == claim_id,
            ParsedField.field_name == "patient_name",
        )
        .first()
    )
    if pf_name and pf_name[0]:
        pf_dob = (
            db.query(ParsedField.field_value)
            .filter(
                ParsedField.claim_id == claim_id,
                ParsedField.field_name == "dob",
            )
            .first()
        )
        return pf_name[0], pf_dob[0] if pf_dob else None

    return None, None


def _upsert_identity_validation(
    db: Session,
    *,
    claim_id: uuid.UUID,
    document_id: uuid.UUID,
    file_name: str,
    status: str,
    patient_match: str,
    patient_name: str | None,
    dob: str | None,
    excluded: bool,
    needs_manual_review: bool,
    reason: str,
    anchor_locked: bool,
) -> None:
    db.query(DocValidation).filter(
        DocValidation.claim_id == claim_id,
        DocValidation.document_id == document_id,
        DocValidation.doc_type == "IDENTITY_GATE",
    ).delete(synchronize_session=False)

    metadata: dict[str, Any] = {
        "phase": "UPLOAD_IDENTITY_GATE",
        "file_name": file_name,
        "identity_dob": dob,
        "excluded_from_pipeline": excluded,
        "needs_manual_review": needs_manual_review,
        "reason": reason,
        "anchor_locked": anchor_locked,
        "checked_at_utc": datetime.utcnow().isoformat() + "Z",
    }

    db.add(DocValidation(
        claim_id=claim_id,
        document_id=document_id,
        status=status,
        doc_type="IDENTITY_GATE",
        doc_type_label="Identity Gate",
        is_medical=1,
        patient_match=patient_match,
        confidence=1.0,
        patient_name=patient_name,
        patient_id_extracted=None,
        issues=[reason],
        validation_metadata=metadata,
    ))


def _apply_identity_gate(
    db: Session,
    claim_id: uuid.UUID,
    documents: list[Document],
) -> dict[str, Any]:
    anchor_name, anchor_dob = _existing_identity_anchor(db, claim_id)
    anchor_name_key = _canonical_name(anchor_name)

    accepted_docs: list[str] = []
    rejected_docs: list[dict[str, str]] = []
    manual_review_required = False

    for doc in documents:
        text = _extract_text_for_identity(Path(doc.minio_path), doc.file_type)
        
        # Check if text is empty or too short (meaning image, scanned PDF, or empty doc)
        if not text or len(text.strip()) < 20:
            # Synchronous text extraction was not possible or returned minimal text.
            # Accept it for the pipeline so it can be OCR'd and validated asynchronously.
            _upsert_identity_validation(
                db,
                claim_id=claim_id,
                document_id=doc.id,
                file_name=doc.file_name,
                status="VALID",
                patient_match="PENDING",
                patient_name=None,
                dob=None,
                excluded=False,
                needs_manual_review=False,
                reason="Document requires OCR for identity verification",
                anchor_locked=False,
            )
            accepted_docs.append(doc.file_name)
            continue

        # Check if this document is actually an identity proof
        is_identity_doc = any(kw in text.lower() for kw in ("government of india", "unique identification authority", "uidai", "income tax department", "permanent account number", "voter id", "passport", "driving licence", "identity card"))
        if not is_identity_doc:
            _upsert_identity_validation(
                db,
                claim_id=claim_id,
                document_id=doc.id,
                file_name=doc.file_name,
                status="VALID",
                patient_match="SKIP",
                patient_name=None,
                dob=None,
                excluded=False,
                needs_manual_review=False,
                reason="Medical or non-KYC document bypassed Identity Gate",
                anchor_locked=False,
            )
            accepted_docs.append(doc.file_name)
            continue

        patient_name, dob_raw = _extract_identity_from_text(text)
        dob = _normalize_dob(dob_raw) if dob_raw else ""

        if not patient_name:
            manual_review_required = True
            reason = "Document missing required patient_name"
            _upsert_identity_validation(
                db,
                claim_id=claim_id,
                document_id=doc.id,
                file_name=doc.file_name,
                status="INVALID",
                patient_match="NO_DATA",
                patient_name=patient_name,
                dob=dob_raw,
                excluded=True,
                needs_manual_review=True,
                reason=reason,
                anchor_locked=False,
            )
            rejected_docs.append({"file_name": doc.file_name, "reason": reason})
            continue

        name_key = _canonical_name(patient_name)

        if not anchor_name_key:
            anchor_name = patient_name
            anchor_dob = dob
            anchor_name_key = name_key
            _upsert_identity_validation(
                db,
                claim_id=claim_id,
                document_id=doc.id,
                file_name=doc.file_name,
                status="VALID",
                patient_match="MATCH",
                patient_name=patient_name,
                dob=dob,
                excluded=False,
                needs_manual_review=False,
                reason="Anchor identity established (name-only)",
                anchor_locked=True,
            )
            accepted_docs.append(doc.file_name)
            continue

        if name_key == anchor_name_key:
            _upsert_identity_validation(
                db,
                claim_id=claim_id,
                document_id=doc.id,
                file_name=doc.file_name,
                status="VALID",
                patient_match="MATCH",
                patient_name=patient_name,
                dob=dob,
                excluded=False,
                needs_manual_review=False,
                reason="Identity matched claim anchor (name-only)",
                anchor_locked=False,
            )
            accepted_docs.append(doc.file_name)
            continue

        manual_review_required = True
        reason = "Patient name mismatch with first-batch claim anchor"
        _upsert_identity_validation(
            db,
            claim_id=claim_id,
            document_id=doc.id,
            file_name=doc.file_name,
            status="INVALID",
            patient_match="MISMATCH",
            patient_name=patient_name,
            dob=dob,
            excluded=True,
            needs_manual_review=True,
            reason=reason,
            anchor_locked=False,
        )
        rejected_docs.append({"file_name": doc.file_name, "reason": reason})

    return {
        "accepted_count": len(accepted_docs),
        "accepted_docs": accepted_docs,
        "rejected_docs": rejected_docs,
        "manual_review_required": manual_review_required,
        "anchor_name": anchor_name,
        "anchor_dob": anchor_dob,
    }


# ------------------------------------------------------------------ routes
router = APIRouter()


def _ensure_users_password_hash_column() -> None:
    with engine.begin() as connection:
        connection.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS password_hash TEXT"))


class RegisterUserIn(BaseModel):
    username: str
    password: str | None = None
    password_hash: str | None = None
    role: str
    first_name: str | None = None
    last_name: str | None = None
    phone: str | None = None
    organization: str | None = None
    employee_id: str | None = None
    dob: str | None = None
    gender: str | None = None
    policy: str | None = None
    sum_insured: Any | None = None
    provider: str | None = "local"


class LoginUserIn(BaseModel):
    username: str
    password: str | None = None
    password_hash: str | None = None
    role: str


@router.post("/auth/register", status_code=201)
def register_local_user(payload: RegisterUserIn):
    """Register a user (patient or TPA) directly in the database without Keycloak/Entra requirement."""
    _ensure_users_password_hash_column()

    email = str(payload.username).strip().lower()
    if not email:
        raise HTTPException(status_code=400, detail="Username/Email is required")

    role_str = str(payload.role).lower()
    normalized_role = "reviewer" if role_str in ("tpa", "reviewer") else "submitter"

    first_name = (payload.first_name or email.split("@")[0] or "User").strip()
    last_name = (payload.last_name or "").strip()

    # Parse sum_insured if provided
    sum_insured_val = None
    if payload.sum_insured is not None and str(payload.sum_insured).strip() != "":
        try:
            sum_insured_val = float(payload.sum_insured)
        except (ValueError, TypeError):
            sum_insured_val = None

    # Parse dob if provided
    dob_val = None
    if payload.dob and str(payload.dob).strip() != "":
        try:
            dob_val = datetime.strptime(str(payload.dob).strip(), "%Y-%m-%d").date()
        except (ValueError, TypeError):
            dob_val = None

    with SessionLocal() as db:
        try:
            # 1. Create or get User
            user_row = db.execute(
                text("SELECT id FROM users WHERE lower(email) = lower(:email)"),
                {"email": email},
            ).mappings().first()

            supplied_hash = (payload.password_hash or "").strip()
            password_hash = supplied_hash or (hash_password(payload.password) if payload.password else None)

            if user_row:
                user_id = user_row["id"]
                db.execute(
                    text("""
                        UPDATE users
                        SET status = 'ACTIVE',
                            password_hash = COALESCE(:password_hash, password_hash),
                            updated_at = now()
                        WHERE id = :id
                    """),
                    {"id": user_id, "password_hash": password_hash},
                )
            else:
                user_row = db.execute(
                    text("""
                        INSERT INTO users (email, phone, external_provider, external_subject_id, status, email_verified, password_hash)
                        VALUES (:email, :phone, 'local', :email, 'ACTIVE', true, :password_hash)
                        RETURNING id
                    """),
                    {"email": email, "phone": payload.phone or None, "password_hash": password_hash},
                ).mappings().one()
                user_id = user_row["id"]

            # 2. Assign Role
            role_row = db.execute(
                text("SELECT id FROM roles WHERE name = :role_name"),
                {"role_name": normalized_role},
            ).mappings().first()

            if not role_row:
                role_row = db.execute(
                    text("INSERT INTO roles (name, description) VALUES (:name, :desc) RETURNING id"),
                    {"name": normalized_role, "desc": f"{normalized_role.title()} role"},
                ).mappings().one()

            role_id = role_row["id"]

            db.execute(
                text("""
                    INSERT INTO user_roles (user_id, role_id)
                    VALUES (:user_id, :role_id)
                    ON CONFLICT (user_id, role_id) DO NOTHING
                """),
                {"user_id": user_id, "role_id": role_id},
            )

            # 3. Create/Update Profile depending on role
            if normalized_role == "submitter":
                # Patient profile
                existing_profile = db.execute(
                    text("SELECT id FROM patient_profiles WHERE user_id = :user_id"),
                    {"user_id": user_id},
                ).mappings().first()

                if existing_profile:
                    db.execute(
                        text("""
                            UPDATE patient_profiles
                            SET first_name = :first_name,
                                last_name = :last_name,
                                dob = COALESCE(:dob, dob),
                                gender = COALESCE(:gender, gender),
                                policy_number = COALESCE(:policy, policy_number),
                                sum_insured = COALESCE(:sum_insured, sum_insured),
                                updated_at = now()
                            WHERE user_id = :user_id
                        """),
                        {
                            "user_id": user_id,
                            "first_name": first_name,
                            "last_name": last_name,
                            "dob": dob_val,
                            "gender": payload.gender or None,
                            "policy": payload.policy or None,
                            "sum_insured": sum_insured_val,
                        },
                    )
                else:
                    db.execute(
                        text("""
                            INSERT INTO patient_profiles (user_id, first_name, last_name, dob, gender, policy_number, sum_insured)
                            VALUES (:user_id, :first_name, :last_name, :dob, :gender, :policy, :sum_insured)
                        """),
                        {
                            "user_id": user_id,
                            "first_name": first_name,
                            "last_name": last_name,
                            "dob": dob_val,
                            "gender": payload.gender or None,
                            "policy": payload.policy or None,
                            "sum_insured": sum_insured_val,
                        },
                    )
            else:
                # Reviewer / TPA staff profile
                org_name = (payload.organization or "Default TPA").strip()
                org_row = db.execute(
                    text("SELECT id FROM organizations WHERE lower(name) = lower(:name) AND type = 'TPA'"),
                    {"name": org_name},
                ).mappings().first()

                if not org_row:
                    org_row = db.execute(
                        text("""
                            INSERT INTO organizations (name, type, status)
                            VALUES (:name, 'TPA', 'ACTIVE')
                            RETURNING id
                        """),
                        {"name": org_name},
                    ).mappings().one()

                org_id = org_row["id"]

                existing_staff = db.execute(
                    text("SELECT id FROM staff_profiles WHERE user_id = :user_id"),
                    {"user_id": user_id},
                ).mappings().first()

                if existing_staff:
                    db.execute(
                        text("""
                            UPDATE staff_profiles
                            SET first_name = :first_name,
                                last_name = :last_name,
                                organization_id = :org_id,
                                employee_id = COALESCE(:employee_id, employee_id),
                                updated_at = now()
                            WHERE user_id = :user_id
                        """),
                        {
                            "user_id": user_id,
                            "first_name": first_name,
                            "last_name": last_name,
                            "org_id": org_id,
                            "employee_id": payload.employee_id or None,
                        },
                    )
                else:
                    db.execute(
                        text("""
                            INSERT INTO staff_profiles (user_id, organization_id, first_name, last_name, employee_id, designation, status)
                            VALUES (:user_id, :org_id, :first_name, :last_name, :employee_id, 'TPA Reviewer', 'ACTIVE')
                        """),
                        {
                            "user_id": user_id,
                            "org_id": org_id,
                            "first_name": first_name,
                            "last_name": last_name,
                            "employee_id": payload.employee_id or None,
                        },
                    )

            db.commit()
            return {
                "success": True,
                "user_id": str(user_id),
                "email": email,
                "role": normalized_role,
                "message": "User registered and profile stored in database successfully",
            }
        except Exception as exc:
            db.rollback()
            logger.exception("Failed to register user locally in database")
            raise HTTPException(status_code=500, detail=f"Database registration error: {str(exc)}") from exc



@router.post("/auth/login", status_code=200)
def login_local_user(payload: LoginUserIn):
    """Authenticate a locally registered user using a stored password hash."""
    _ensure_users_password_hash_column()

    email = str(payload.username).strip().lower()
    role_str = str(payload.role).strip().lower()
    if not email or not role_str or (not payload.password and not payload.password_hash):
        raise HTTPException(status_code=400, detail="Username, password, and role are required")

    normalized_role = "reviewer" if role_str in ("tpa", "reviewer") else "submitter"

    with force_master_session():
        with SessionLocal() as db:
            user_row = db.execute(
                text("SELECT id, password_hash, status FROM users WHERE lower(email) = lower(:email)"),
                {"email": email},
            ).mappings().first()

            if not user_row:
                raise HTTPException(status_code=401, detail="Username not found")

            if user_row["status"] in ("BLOCKED", "DELETED"):
                raise HTTPException(status_code=403, detail="Account is not active")

            supplied_hash = (payload.password_hash or "").strip()
            supplied_password = payload.password or ""
            stored_hash = user_row["password_hash"]
            if not stored_hash:
                raise HTTPException(status_code=401, detail="Invalid email or password")

            if supplied_hash:
                if not password_matches(supplied_hash, stored_hash):
                    raise HTTPException(status_code=401, detail="Invalid password")
            elif not password_matches(supplied_password, stored_hash):
                raise HTTPException(status_code=401, detail="Invalid password")

            role_match = db.execute(
                text(
                    "SELECT 1 FROM user_roles ur JOIN roles r ON ur.role_id = r.id WHERE ur.user_id = :user_id AND r.name = :role_name"
                ),
                {"user_id": user_row["id"], "role_name": normalized_role},
            ).scalar()

            if not role_match:
                actual_role_row = db.execute(
                    text(
                        "SELECT r.name FROM user_roles ur JOIN roles r ON ur.role_id = r.id WHERE ur.user_id = :user_id ORDER BY r.name LIMIT 1"
                    ),
                    {"user_id": user_row["id"]},
                ).mappings().first()
                actual_role = actual_role_row["name"] if actual_role_row else normalized_role
                raise HTTPException(
                    status_code=403,
                    detail={
                        "message": f"User is not registered as a {role_str}",
                        "actual_role": actual_role,
                    },
                )

            db.execute(
                text("UPDATE users SET last_login_at = now(), updated_at = now() WHERE id = :id"),
                {"id": user_row["id"]},
            )
            db.commit()

    return {
        "success": True,
        "user_id": str(user_row["id"]),
        "email": email,
        "role": normalized_role,
        "message": "Login successful",
    }


# ------------------------------------------------------------------ TPA registration
# Passwords are sent only to Keycloak and are never persisted in ClaimGPT.
class TpaRegistrationIn(BaseModel):
    first_name: str = Field(min_length=1, max_length=100)
    last_name: str = Field(min_length=1, max_length=100)
    email: EmailStr
    phone: str | None = Field(default=None, max_length=30)
    organization_id: uuid.UUID
    employee_id: str | None = Field(default=None, max_length=100)
    password: str = Field(min_length=8, max_length=256)


class OrganizationRegistrationIn(BaseModel):
    name: str = Field(min_length=2, max_length=255)
    address: str | None = Field(default=None, max_length=2000)


def _keycloak_admin_token() -> str:
    """Obtain a short-lived admin token for provisioning a Keycloak user."""
    import httpx

    url = f"{settings.keycloak_url}/realms/master/protocol/openid-connect/token"
    try:
        response = httpx.post(
            url,
            data={
                "grant_type": "password",
                "client_id": "admin-cli",
                "username": settings.keycloak_admin_username,
                "password": settings.keycloak_admin_password,
            },
            timeout=10,
        )
        response.raise_for_status()
        return response.json()["access_token"]
    except Exception as exc:
        logger.exception("Could not authenticate with Keycloak admin API")
        raise HTTPException(status_code=503, detail="Account service is unavailable") from exc


def _create_keycloak_reviewer(payload: TpaRegistrationIn) -> str:
    """Create the credential record in Keycloak and assign the reviewer role."""
    import httpx

    token = _keycloak_admin_token()
    base_url = f"{settings.keycloak_url}/admin/realms/{settings.keycloak_realm}"
    headers = {"Authorization": f"Bearer {token}"}
    user = {
        "username": str(payload.email),
        "email": str(payload.email),
        "firstName": payload.first_name.strip(),
        "lastName": payload.last_name.strip(),
        "enabled": True,
        "emailVerified": False,
        "credentials": [{"type": "password", "value": payload.password, "temporary": False}],
    }
    try:
        response = httpx.post(f"{base_url}/users", headers=headers, json=user, timeout=10)
        if response.status_code == 409:
            raise HTTPException(status_code=409, detail="An account with this email already exists")
        response.raise_for_status()
        location = response.headers.get("Location", "")
        keycloak_user_id = location.rstrip("/").split("/")[-1]
        if not keycloak_user_id:
            raise RuntimeError("Keycloak did not return a user ID")

        role_response = httpx.get(f"{base_url}/roles/reviewer", headers=headers, timeout=10)
        role_response.raise_for_status()
        assignment = httpx.post(
            f"{base_url}/users/{keycloak_user_id}/role-mappings/realm",
            headers=headers,
            json=[role_response.json()],
            timeout=10,
        )
        assignment.raise_for_status()
        return keycloak_user_id
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("Could not provision Keycloak reviewer")
        raise HTTPException(status_code=503, detail="Could not create the login account") from exc


@router.get("/organizations", status_code=200)
def list_approved_tpa_organizations():
    """Public registration list: only organizations already approved by an admin."""
    with SessionLocal() as db:
        rows = db.execute(
            text("SELECT id, name FROM organizations WHERE type = 'TPA' AND status = 'ACTIVE' ORDER BY name")
        ).mappings().all()
    return {"organizations": [{"id": str(row["id"]), "name": row["name"]} for row in rows]}


@router.post("/organizations/registration", status_code=201)
def request_tpa_organization_registration(payload: OrganizationRegistrationIn):
    """Submit an organization for approval; it is intentionally not selectable yet."""
    with SessionLocal() as db:
        existing = db.execute(
            text("SELECT id, status FROM organizations WHERE lower(name) = lower(:name) AND type = 'TPA'"),
            {"name": payload.name.strip()},
        ).mappings().first()
        if existing:
            return {"id": str(existing["id"]), "status": existing["status"], "message": "Organization already exists"}

        row = db.execute(
            text("""
                INSERT INTO organizations (name, type, address, status)
                VALUES (:name, 'TPA', :address, 'PENDING')
                RETURNING id, status
            """),
            {"name": payload.name.strip(), "address": payload.address.strip() if payload.address else None},
        ).mappings().one()
        db.commit()
    return {"id": str(row["id"]), "status": row["status"], "message": "Organization submitted for approval"}


@router.post("/tpa-adjusters/registration", status_code=201)
def register_tpa_adjuster(payload: TpaRegistrationIn):
    """Create a Keycloak login and its matching local staff profile."""
    _ensure_users_password_hash_column()

    with SessionLocal() as db:
        organization = db.execute(
            text("SELECT id FROM organizations WHERE id = :id AND type = 'TPA' AND status = 'ACTIVE'"),
            {"id": str(payload.organization_id)},
        ).mappings().first()
        if not organization:
            raise HTTPException(status_code=400, detail="Select an approved TPA organization")

    keycloak_user_id = _create_keycloak_reviewer(payload)
    try:
        with SessionLocal() as db:
            user = db.execute(
                text("""
                    INSERT INTO users (email, phone, external_provider, external_subject_id, status, password_hash)
                    VALUES (:email, :phone, 'keycloak', :subject_id, 'ACTIVE', :password_hash)
                    RETURNING id
                """),
                {
                    "email": str(payload.email),
                    "phone": payload.phone or None,
                    "subject_id": keycloak_user_id,
                    "password_hash": hash_password(payload.password),
                },
            ).mappings().one()
            db.execute(
                text("""
                    INSERT INTO staff_profiles (user_id, organization_id, employee_id, designation, status)
                    VALUES (:user_id, :organization_id, :employee_id, 'TPA Adjuster', 'ACTIVE')
                """),
                {"user_id": str(user["id"]), "organization_id": str(payload.organization_id), "employee_id": payload.employee_id or None},
            )
            db.execute(
                text("""
                    INSERT INTO user_roles (user_id, role_id)
                    SELECT :user_id, id FROM roles WHERE name = 'reviewer'
                """),
                {"user_id": str(user["id"])},
            )
            db.commit()
    except Exception as exc:
        logger.exception("Keycloak user %s was created but local profile creation failed", keycloak_user_id)
        raise HTTPException(status_code=500, detail="Could not save the account profile") from exc

    return {"user_id": str(user["id"]), "role": "reviewer", "message": "TPA adjuster account created"}


@router.get("/health")
def health():
    db_ok = check_db_health()
    status = "ok" if db_ok else "degraded"
    return {"status": status, "database": "up" if db_ok else "down"}


@router.post("/auth/login")
@router.post("/auth/register")
def authenticate_or_register_user(data: dict[str, Any] = Body(...), db: Session = Depends(get_db)):
    username = data.get("username") or data.get("name") or data.get("email", "Swagath")
    if isinstance(username, str) and "@" in username:
        username = username.split("@")[0]
    username = str(username).strip().capitalize()
    
    email = data.get("email") or f"{username.lower()}@example.com"
    role = data.get("role", "patient")
    
    _audit(db, "USER_LOGIN_OR_REGISTER", metadata={"username": username, "email": email, "role": role})
    logger.info("User registered/authenticated in Docker backend: %s (%s)", username, email)
    return {
        "status": "success",
        "message": f"Account {username} initialized in backend",
        "user": {
            "name": username,
            "email": email,
            "role": role,
            "account_id": f"ACC-{username.upper()}-2026"
        }
    }


@router.post("/claims", status_code=202)
async def create_claim(
    files: list[UploadFile] = File(...),
    policy_id: str = Form(None),
    patient_id: str = Form(None),
):
    """Create a new claim by uploading files.
    
    This endpoint accepts files, saves them to disk, and enqueues the pipeline.
    All database operations (idempotency, deduplication) are handled by the 
    intake_task in the Celery worker.
    """
    logger.info(f"[create_claim] Starting with {len(files)} files")
    upload_log.info(
        "UPLOAD_START | endpoint=create_claim files=%d policy_id=%s patient_id=%s names=%s",
        len(files),
        policy_id,
        patient_id,
        [getattr(f, "filename", "?") for f in files],
    )
    
    if not files:
        upload_log.warning("UPLOAD_REJECTED | endpoint=create_claim reason=no_files")
        raise HTTPException(status_code=400, detail="At least one file is required")

    # --- Validate all files and read content ---
    file_metadata_list: list[dict[str, str]] = []  # Will hold metadata for intake_task
    saved_paths: list[Path] = []
    
    try:
        for idx, file in enumerate(files):
            # Validate content type
            effective_ct, ok = _resolve_content_type(file)
            if not ok:
                upload_log.warning(
                    "UPLOAD_REJECTED | endpoint=create_claim reason=unsupported_type file=%s type=%s",
                    file.filename, file.content_type,
                )
                raise HTTPException(
                    status_code=415,
                    detail=f"Unsupported file type '{file.content_type}' for '{file.filename}'. "
                    f"Allowed: {', '.join(sorted(settings.allowed_content_types))}",
                )
            
            # Read and validate file size
            file_bytes = await file.read()
            if len(file_bytes) > settings.max_upload_bytes:
                upload_log.warning(
                    "UPLOAD_REJECTED | endpoint=create_claim reason=too_large file=%s bytes=%d max=%d",
                    file.filename, len(file_bytes), settings.max_upload_bytes,
                )
                raise HTTPException(
                    status_code=413,
                    detail=f"File '{file.filename}' too large ({len(file_bytes)} bytes). Max: {settings.max_upload_bytes} bytes",
                )
            
            # Calculate content hash and safe filename
            safe_name = _safe_filename(file.filename)
            content_hash = hashlib.sha256(file_bytes).hexdigest()
            logger.info(f"[create_claim] File validated: {safe_name}, hash={content_hash}")
            
            # Generate temporary file name (will be replaced by claim_id once created in intake_task)
            temp_name = f"pending_{uuid.uuid4().hex[:8]}_{safe_name}"
            local_path = RAW_STORAGE / temp_name
            
            # Save file to disk
            try:
                async with aiofiles.open(local_path, "wb") as f:
                    await f.write(file_bytes)
                with open(local_path, "rb+") as sync_f:
                    sync_f.flush()
                    os.fsync(sync_f.fileno())
                saved_paths.append(local_path)
                logger.info(f"[create_claim] File saved: {local_path}")
            except OSError as e:
                logger.exception(f"[create_claim] Failed to write file: {local_path}")
                raise HTTPException(status_code=500, detail="Failed to store uploaded file")
            
            # Store metadata for intake_task
            file_metadata_list.append({
                "path": str(local_path),
                "safe_name": safe_name,
                "content_hash": content_hash,
                "effective_ct": effective_ct,
            })
            
            upload_log.info(
                "FILE_RECEIVED | endpoint=create_claim file=%s bytes=%d type=%s sha256=%s",
                safe_name, len(file_bytes), effective_ct, content_hash,
            )
        
        # --- Enqueue pipeline with file metadata ---
        # The intake_task will:
        # 1. Create the claim in the database
        # 2. Create document rows
        # 3. Check for idempotency/deduplication
        # 4. Move files to permanent location with claim_id
        task_id = _enqueue_pipeline(file_metadata_list, policy_id, patient_id)
        
        upload_log.info(
            "UPLOAD_SUCCESS | endpoint=create_claim files=%d task_id=%s",
            len(file_metadata_list), task_id,
        )
        
        return {
            "task_id": task_id,
            "status": "QUEUED",
            "message": "Claim upload queued. Check status via /claims/{claim_id}/progress endpoint.",
        }
    
    except HTTPException:
        # Clean up saved files on validation error
        for p in saved_paths:
            p.unlink(missing_ok=True)
        raise
    except Exception as exc:
        logger.exception("Error during file upload processing")
        upload_log.exception(
            "UPLOAD_FAILURE | endpoint=create_claim files=%d error=%s",
            len(files), exc,
        )
        # Clean up saved files
        for p in saved_paths:
            p.unlink(missing_ok=True)
        raise HTTPException(status_code=500, detail="Failed to process file upload")


@router.get("/claims", response_model=ClaimListOut)
def list_claims(
    offset: int = Query(0, ge=0),
    limit: int = Query(20, ge=1, le=100),
    db: Session = Depends(get_db),
):
    try:
        total = db.query(Claim).count()
        claims = (
            db.query(Claim)
            .order_by(Claim.created_at.desc())
            .offset(offset)
            .limit(limit)
            .all()
        )
        
        # Batch-fetch relevant parsed fields for these claims to avoid N+1 queries
        if claims:
            claim_ids = [c.id for c in claims]
            pf_rows = db.query(ParsedField).filter(
                ParsedField.claim_id.in_(claim_ids),
                ParsedField.field_name.in_([
                    "patient_name", "member_name", "insured_name",
                    "hospital_name", "hospital",
                    "doctor_name", "doctor", "provider_name", "rendering_provider",
                    "diagnosis", "primary_diagnosis", "chief_complaint"
                ])
            ).all()
            
            from collections import defaultdict
            pf_by_claim = defaultdict(dict)
            for row in pf_rows:
                pf_by_claim[row.claim_id][row.field_name] = row.field_value
                
            for c in claims:
                fields = pf_by_claim[c.id]
                c.patient_name = fields.get("patient_name") or fields.get("member_name") or fields.get("insured_name") or None
                c.hospital_name = fields.get("hospital_name") or fields.get("hospital") or None
                c.doctor_name = fields.get("doctor_name") or fields.get("doctor") or fields.get("provider_name") or fields.get("rendering_provider") or None
                c.diagnosis = fields.get("diagnosis") or fields.get("primary_diagnosis") or fields.get("chief_complaint") or None

        claim_items = [
            {
                "id": c.id,
                "policy_id": c.policy_id,
                "patient_id": c.patient_id,
                "status": c.status,
                "source": c.source,
                "created_at": c.created_at,
                "updated_at": c.updated_at,
                "documents": c.documents,
                "task_id": getattr(c, "task_id", None),
                "patient_name": getattr(c, "patient_name", None),
                "hospital_name": getattr(c, "hospital_name", None),
                "doctor_name": getattr(c, "doctor_name", None),
                "diagnosis": getattr(c, "diagnosis", None),
            }
            for c in claims
        ]

        return ClaimListOut(claims=claim_items, total=total)
    except Exception as exc:
        logger.exception("Error listing claims")
        raise HTTPException(status_code=500, detail=f"Failed to list claims: {str(exc)}")


@router.get("/claims/{claim_id}", response_model=ClaimOut)
def get_claim(claim_id: str, db: Session = Depends(get_db)):
    cid = _parse_uuid(claim_id)
    claim = (
        db.query(Claim)
        .options(selectinload(Claim.documents))
        .filter(Claim.id == cid)
        .first()
    )
    if not claim:
        raise HTTPException(status_code=404, detail="Claim not found")
        
    # Fetch relevant parsed fields for this claim
    pf_rows = db.query(ParsedField).filter(
        ParsedField.claim_id == cid,
        ParsedField.field_name.in_([
            "patient_name", "member_name", "insured_name",
            "hospital_name", "hospital",
            "doctor_name", "doctor", "provider_name", "rendering_provider",
            "diagnosis", "primary_diagnosis", "chief_complaint"
        ])
    ).all()
    
    fields = {row.field_name: row.field_value for row in pf_rows}
    claim.patient_name = fields.get("patient_name") or fields.get("member_name") or fields.get("insured_name") or None
    claim.hospital_name = fields.get("hospital_name") or fields.get("hospital") or None
    claim.doctor_name = fields.get("doctor_name") or fields.get("doctor") or fields.get("provider_name") or fields.get("rendering_provider") or None
    claim.diagnosis = fields.get("diagnosis") or fields.get("primary_diagnosis") or fields.get("chief_complaint") or None
    
    return ClaimOut.model_validate(claim).model_dump(mode="json")


def _map_progress(current_step: str | None, status: str | None) -> tuple[str | None, int]:
    if current_step == "STARTING":
        return "Starting", 5
    if current_step == "OCR_IN_PROGRESS":
        return "OCR (extracting text)", 20
    if current_step == "OCR_COMPLETED":
        return "OCR complete", 35
    if current_step == "PARSING_IN_PROGRESS":
        return "Parsing (LLM agent reading document)", 55
    if current_step == "PARSING_COMPLETED":
        return "Parsing complete", 70
    if current_step == "CODING_ANALYSIS":
        return "Medical coding (ICD-10 / CPT)", 78
    if current_step == "CODING_COMPLETED":
        return "Coding complete", 82
    if current_step == "RISK_ANALYSIS":
        return "Risk scoring", 86
    if current_step == "RISK_COMPLETED":
        return "Risk complete", 90
    if current_step == "VALIDATION_RUNNING":
        return "Validating", 92
    if current_step == "VALIDATION_COMPLETED":
        return "Validation complete", 96
    if current_step == "RETRYING":
        # Don't regress — keep above prior steps; monotonic guard below also protects.
        return "Retrying (transient)", 92
    if current_step == "FAILED" or status == "FAILED":
        return "Failed", 0
    if current_step == "FINALIZING":
        return "Finalizing", 98
    if current_step == "FINISHED" or status == "FINISHED":
        return "Completed", 100
    return current_step, 0


# Progress cache removed to prevent uvicorn multi-worker state mismatch


@router.get("/claims/{claim_id}/status")
def get_claim_status(claim_id: str, db: Session = Depends(get_db)):
    cid = _parse_uuid(claim_id)
    state = get_latest_workflow_state(db, cid)
    if not state:
        return {"current_step": None, "status": None, "step_index": 0, "percentage": 0.0}
    
    step_index = _get_step_index(state.current_step, state.status)
    percentage = (step_index / 5) * 100 if step_index > 0 else 0.0
    return {
        "current_step": state.current_step,
        "status": state.status,
        "step_index": step_index,
        "percentage": percentage
    }


@router.get("/claims/{claim_id}/progress")
def get_claim_progress(claim_id: str, db: Session = Depends(get_db)):
    cid = _parse_uuid(claim_id)
    state = get_latest_workflow_state(db, cid)

    # No workflow state yet: distinguish "claim does not exist" from
    # "claim was created but the pipeline hasn't recorded any progress yet".
    if not state:
        claim = db.query(Claim).filter(Claim.id == cid).first()
        if not claim:
            raise HTTPException(status_code=404, detail="Claim not found")
        # Claim exists but no state row — treat as queued, never as silently null.
        return {
            "status": "QUEUED",
            "step": "Queued (waiting for worker)",
            "percentage": 2,
            "is_complete": False,
            "error": None,
        }

    step, percentage = _map_progress(state.current_step, state.status)
    is_failed = (state.status == "FAILED") or (state.current_step == "FAILED")
    is_complete = bool(percentage == 100 or is_failed)

    error_message: str | None = None
    if is_failed:
        # Surface the most recent job error message so the UI can show *why*
        # the upload stopped, instead of polling forever on 0%.
        try:
            latest_parse = (
                db.query(ParseJob)
                .filter(ParseJob.claim_id == cid)
                .order_by(ParseJob.created_at.desc())
                .first()
            )
            if latest_parse and latest_parse.error_message:
                error_message = latest_parse.error_message
            if not error_message:
                from libs.shared.models import OcrJob as _OcrJob
                latest_ocr = (
                    db.query(_OcrJob)
                    .filter(_OcrJob.claim_id == cid)
                    .order_by(_OcrJob.created_at.desc())
                    .first()
                )
                if latest_ocr and latest_ocr.error_message:
                    error_message = latest_ocr.error_message
        except Exception:
            logger.exception("Failed to read latest job error for claim %s", cid)
        if not error_message:
            error_message = "Pipeline failed. See server logs for details."

    # Mapped from database state directly, naturally monotonic in Celery chain
    return {
        "status": state.status,
        "step": step,
        "percentage": percentage,
        "is_complete": is_complete,
        "error": error_message,
    }


@router.get("/claims/{claim_id}/file")
def download_original_file(claim_id: str, view: bool = False, db: Session = Depends(get_db)):
    cid = _parse_uuid(claim_id)

    doc = (
        db.query(Document)
        .filter(Document.claim_id == cid)
        .order_by(Document.uploaded_at.desc())
        .first()
    )
    if not doc:
        raise HTTPException(status_code=404, detail="No document found for claim")

    disp = "inline" if view else "attachment"

    if doc.minio_path and doc.minio_path.startswith("s3://"):
        from libs.shared.storage import MinioStorage
        from fastapi.responses import StreamingResponse
        client = MinioStorage.get_client()
        bucket = MinioStorage.BUCKET_NAME
        s3_key = doc.minio_path[len(f"s3://{bucket}/"):]
        try:
            response = client.get_object(Bucket=bucket, Key=s3_key)
            return StreamingResponse(
                response["Body"].iter_chunks(),
                media_type=doc.file_type or "application/octet-stream",
                headers={
                    "Content-Disposition": f'{disp}; filename="{doc.file_name}"'
                }
            )
        except Exception as e:
            logger.exception(f"Failed to fetch {doc.minio_path} from S3: {e}")
            raise HTTPException(status_code=404, detail="File not found in cloud storage")

    file_path = Path(doc.minio_path).resolve()

    # prevent path traversal — file must be under RAW_STORAGE
    if not str(file_path).startswith(str(RAW_STORAGE)):
        raise HTTPException(status_code=403, detail="Access denied")

    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Stored file not found on disk")

    return FileResponse(
        str(file_path),
        media_type=doc.file_type or "application/pdf",
        headers={"Content-Disposition": f'{disp}; filename="{doc.file_name}"'}
    )


@router.get("/claims/{claim_id}/documents/{doc_id}/file")
def download_document_file(claim_id: str, doc_id: str, view: bool = False, db: Session = Depends(get_db)):
    cid = _parse_uuid(claim_id)
    did = _parse_uuid(doc_id)

    doc = db.query(Document).filter(Document.id == did, Document.claim_id == cid).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    disp = "inline" if view else "attachment"

    if doc.minio_path and doc.minio_path.startswith("s3://"):
        from libs.shared.storage import MinioStorage
        from fastapi.responses import StreamingResponse
        client = MinioStorage.get_client()
        bucket = MinioStorage.BUCKET_NAME
        s3_key = doc.minio_path[len(f"s3://{bucket}/"):]
        try:
            response = client.get_object(Bucket=bucket, Key=s3_key)
            return StreamingResponse(
                response["Body"].iter_chunks(),
                media_type=doc.file_type or "application/octet-stream",
                headers={
                    "Content-Disposition": f'{disp}; filename="{doc.file_name}"'
                }
            )
        except Exception as e:
            logger.exception(f"Failed to fetch {doc.minio_path} from S3: {e}")
            raise HTTPException(status_code=404, detail="File not found in cloud storage")

    file_path = Path(doc.minio_path).resolve()

    if not file_path.exists():
        file_path = (RAW_STORAGE / Path(doc.minio_path).name).resolve()
        if not file_path.exists():
            raise HTTPException(status_code=404, detail="Original document file missing on disk")

    # prevent path traversal — file must be under RAW_STORAGE
    if not str(file_path).startswith(str(RAW_STORAGE)):
        raise HTTPException(status_code=403, detail="Access denied")

    return FileResponse(
        str(file_path),
        media_type=doc.file_type or "application/pdf",
        headers={"Content-Disposition": f'{disp}; filename="{doc.file_name}"'}
    )


@router.get("/claims/{claim_id}/documents/{doc_id}/pages/{page_number}/image")
def get_document_page_image(claim_id: str, doc_id: str, page_number: int = 1, db: Session = Depends(get_db)):
    cid = _parse_uuid(claim_id)
    did = _parse_uuid(doc_id)

    doc = db.query(Document).filter(Document.id == did, Document.claim_id == cid).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    from fastapi.responses import Response

    if doc.minio_path and doc.minio_path.startswith("s3://"):
        from libs.shared.storage import MinioStorage
        client = MinioStorage.get_client()
        bucket = MinioStorage.BUCKET_NAME
        s3_key = doc.minio_path[len(f"s3://{bucket}/"):]
        try:
            response = client.get_object(Bucket=bucket, Key=s3_key)
            file_bytes = response["Body"].read()
        except Exception as e:
            logger.exception(f"Failed to fetch {doc.minio_path} from S3: {e}")
            raise HTTPException(status_code=404, detail="File not found in cloud storage")

        # Determine file extension from key/filename
        ext = Path(doc.file_name).suffix.lower() if doc.file_name else ".pdf"
        if ext in {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif"} or (doc.file_type and doc.file_type.startswith("image/")):
            return Response(content=file_bytes, media_type=doc.file_type or "image/png")

        if ext == ".pdf" or doc.file_type == "application/pdf":
            try:
                import io
                import pypdfium2
                pdf = pypdfium2.PdfDocument(file_bytes)
                total_pages = len(pdf)
                target_idx = max(0, min(page_number - 1, total_pages - 1))
                page = pdf.get_page(target_idx)
                img = page.render(scale=2).to_pil()
                buf = io.BytesIO()
                img.save(buf, format="PNG")
                buf.seek(0)
                return Response(content=buf.getvalue(), media_type="image/png")
            except Exception as e:
                logger.exception("Failed to render PDF page image from S3: %s", e)
                raise HTTPException(status_code=500, detail="Failed to render PDF page")

        return Response(content=file_bytes, media_type=doc.file_type or "application/octet-stream")

    file_path = Path(doc.minio_path).resolve()
    if not file_path.exists():
        file_path = (RAW_STORAGE / Path(doc.minio_path).name).resolve()
        if not file_path.exists():
            raise HTTPException(status_code=404, detail="Original document file missing on disk")

    ext = file_path.suffix.lower()
    if ext in {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif"}:
        return FileResponse(str(file_path), media_type=doc.file_type or "image/png")

    if ext == ".pdf" or doc.file_type == "application/pdf":
        try:
            import io
            import pypdfium2
            pdf = pypdfium2.PdfDocument(str(file_path))
            total_pages = len(pdf)
            target_idx = max(0, min(page_number - 1, total_pages - 1))
            page = pdf.get_page(target_idx)
            img = page.render(scale=2).to_pil()
            buf = io.BytesIO()
            img.save(buf, format="PNG")
            buf.seek(0)
            return Response(content=buf.getvalue(), media_type="image/png")
        except Exception as e:
            logger.exception("Failed to render PDF page image: %s", e)
            raise HTTPException(status_code=500, detail="Failed to render PDF page")

    return FileResponse(str(file_path), media_type=doc.file_type or "application/octet-stream")


@router.post("/claims/{claim_id}/documents", response_model=ClaimOut, status_code=201)
async def add_documents_to_claim(
    claim_id: str,
    files: list[UploadFile] = File(...),
    db: Session = Depends(get_db),
):
    logger.info(f"[IDEMPOTENCY] Starting add_documents_to_claim with {len(files)} files for claim {claim_id}.")
    upload_log.info(
        "UPLOAD_START | endpoint=add_documents claim_id=%s files=%d names=%s",
        claim_id, len(files), [getattr(f, "filename", "?") for f in files],
    )
    """Add supporting documents to an existing claim."""
    cid = _parse_uuid(claim_id)
    claim = db.query(Claim).filter(Claim.id == cid).first()
    if not claim:
        upload_log.warning(
            "UPLOAD_REJECTED | endpoint=add_documents reason=claim_not_found claim_id=%s",
            claim_id,
        )
        raise HTTPException(status_code=404, detail="Claim not found")

    if not files:
        upload_log.warning(
            "UPLOAD_REJECTED | endpoint=add_documents reason=no_files claim_id=%s",
            claim_id,
        )
        raise HTTPException(status_code=400, detail="At least one file is required")

    # --- validate all files and calculate content_hash
    file_data: list[tuple[UploadFile, bytes, str, str, str]] = []  # (file, bytes, safe_name, content_hash, effective_ct)
    for file in files:
        effective_ct, ok = _resolve_content_type(file)
        if not ok:
            raise HTTPException(
                status_code=415,
                detail=f"Unsupported file type '{file.content_type}' for '{file.filename}'. "
                f"Allowed: {', '.join(sorted(settings.allowed_content_types))}",
            )
        file_bytes = await file.read()
        if len(file_bytes) > settings.max_upload_bytes:
            raise HTTPException(
                status_code=413,
                detail=f"File '{file.filename}' too large ({len(file_bytes)} bytes). Max: {settings.max_upload_bytes} bytes",
            )
        safe_name = _safe_filename(file.filename)
        content_hash = hashlib.sha256(file_bytes).hexdigest()
        logger.info(f"[IDEMPOTENCY] Calculated content_hash for file '{safe_name}': {content_hash}")
        file_data.append((file, file_bytes, safe_name, content_hash, effective_ct))


    # --- count existing docs for naming
    existing_count = db.query(Document).filter(Document.claim_id == cid).count()

    # --- save files and create document rows
    saved_paths: list[Path] = []
    new_docs: list[Document] = []
    new_doc_added = False
    for idx, (file, file_bytes, safe_name, content_hash, effective_ct) in enumerate(file_data):
        # --- DUPLICATE CHECK LOGIC ---
        # 1. Calculate SHA-256 hash of file bytes (content_hash)
        # 2. Query Document table for any document with same claim_id and content_hash
        logger.info(f"[IDEMPOTENCY] Checking for duplicate: claim_id={claim.id}, content_hash={content_hash}")
        duplicate_doc = db.query(Document).filter(Document.claim_id == claim.id, Document.content_hash == content_hash).first()
        if duplicate_doc:
            logger.info(f"[IDEMPOTENCY] Duplicate document detected for claim {claim.id} and hash {content_hash}, skipping upload and returning existing document.")
            _audit(db, "DUPLICATE_DOCUMENT_SKIPPED", claim_id=claim.id, metadata={
                "file_name": safe_name,
                "content_hash": content_hash,
                "existing_document_id": str(duplicate_doc.id),
            })
            continue  # skip adding duplicate

        ext = Path(safe_name).suffix or ".bin"
        stored_name = f"{claim.id}_{existing_count + idx}{ext}"
        local_path = RAW_STORAGE / stored_name

        try:
            async with aiofiles.open(local_path, "wb") as f:
                await f.write(file_bytes)
            saved_paths.append(local_path)
        except OSError:
            for p in saved_paths:
                p.unlink(missing_ok=True)
            db.rollback()
            logger.exception("Failed to write uploaded file to disk")
            raise HTTPException(status_code=500, detail="Failed to store uploaded file")

        doc = Document(
            claim_id=claim.id,
            file_name=safe_name,
            file_type=effective_ct,
            minio_path=str(local_path),
            content_hash=content_hash,
        )
        db.add(doc)
        new_docs.append(doc)
        new_doc_added = True

    if not new_doc_added:
        logger.info(f"No new documents added for claim {claim.id}; all uploads were duplicates.")
        _audit(db, "DUPLICATE_DOCUMENTS_ALL_SKIPPED", claim_id=claim.id, metadata={
            "file_count": len(file_data),
            "reason": "All uploaded documents were duplicates. Pipeline will still be triggered to ensure combined report."
        })
        # Always trigger pipeline to ensure combined report
        try:
            claim.status = "UPLOADED"
            db.commit()
            task_id = _enqueue_pipeline(str(claim.id))
        except Exception:
            db.rollback()
            logger.exception("Failed to enqueue Celery pipeline for claim %s", claim.id)
            raise HTTPException(status_code=503, detail="No new documents, but failed to enqueue background tasks for combined report")
        payload = _build_claim_response(db, cid, {"task_id": task_id})
        return JSONResponse(status_code=200, content=payload)

    db.flush()
    gate_result = _apply_identity_gate(db, claim.id, new_docs)
    manual_review_message = None
    if gate_result["accepted_count"] == 0:
        claim.status = "MANUAL_REVIEW_REQUIRED"
        upsert_workflow_state(db, claim.id, "MANUAL_REVIEW_REQUIRED", status="FAILED")
        from libs.shared.storage import MinioStorage
        for doc in new_docs:
            if doc.minio_path:
                try:
                    MinioStorage.delete_file(doc.minio_path)
                except Exception:
                    logger.warning("Failed to delete mismatched file from MinIO: %s", doc.minio_path)
            db.delete(doc)
        manual_review_message = (
            "Manual review required: Patient name mismatch detected in the documents you added. "
            "Please check that the uploaded documents have the correct patient details."
        )
    else:
        claim.status = "UPLOADED"

    try:
        db.commit()
    except Exception as exc:
        db.rollback()
        for p in saved_paths:
            p.unlink(missing_ok=True)
        logger.exception("DB commit failed adding documents")
        upload_log.exception(
            "UPLOAD_FAILURE | endpoint=add_documents claim_id=%s stage=db_commit error=%s",
            claim.id, exc,
        )
        raise HTTPException(status_code=500, detail="Failed to save documents")

    logger.info("Added %d docs to claim %s", len(new_docs), claim.id)
    db.refresh(claim)

    task_id: str | None = None
    if gate_result["accepted_count"] > 0:
        try:
            task_id = _enqueue_pipeline(str(claim.id))
        except Exception as exc:
            logger.exception("Failed to enqueue Celery pipeline for claim %s", claim.id)
            upload_log.exception(
                "UPLOAD_FAILURE | endpoint=add_documents claim_id=%s stage=enqueue_pipeline error=%s",
                claim.id, exc,
            )
            raise HTTPException(status_code=503, detail="Documents saved but failed to enqueue background tasks")
    else:
        logger.warning("Claim %s no accepted new docs after identity gate; workflow not retriggered", claim.id)
        upload_log.warning(
            "UPLOAD_PARTIAL | endpoint=add_documents claim_id=%s reason=identity_gate_rejected_all",
            claim.id,
        )

    extra = {"task_id": task_id} if task_id else {}
    if manual_review_message:
        extra["manual_review_reason"] = manual_review_message
    payload = _build_claim_response(db, cid, extra)
    _audit(db, "DOCUMENTS_ADDED", claim_id=claim.id, metadata={
        "files": [s for _, _, s, _, _ in file_data],
        "file_count": len(new_docs),
        "total_documents": existing_count + len(new_docs),
        "identity_gate": gate_result,
        "manual_review_reason": manual_review_message,
    })
    upload_log.info(
        "UPLOAD_SUCCESS | endpoint=add_documents claim_id=%s new_docs=%d total=%d task_id=%s",
        claim.id, len(new_docs), existing_count + len(new_docs), task_id,
    )
    return payload


@router.delete("/claims/{claim_id}/documents/{doc_id}", response_model=ClaimOut)
def delete_document(
    claim_id: str,
    doc_id: str,
    db: Session = Depends(get_db),
):
    """Delete a single document from a claim."""
    cid = _parse_uuid(claim_id)
    did = _parse_uuid(doc_id)
    claim = db.query(Claim).filter(Claim.id == cid).first()
    if not claim:
        raise HTTPException(status_code=404, detail="Claim not found")

    doc = db.query(Document).filter(Document.id == did, Document.claim_id == cid).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # Prevent deleting the last document
    doc_count = db.query(Document).filter(Document.claim_id == cid).count()
    if doc_count <= 1:
        raise HTTPException(status_code=400, detail="Cannot delete the only document. Delete the claim instead.")

    # Remove file from disk
    try:
        p = Path(doc.minio_path).resolve()
        if str(p).startswith(str(RAW_STORAGE)):
            p.unlink(missing_ok=True)
    except OSError:
        logger.warning("Failed to delete file %s", doc.minio_path)

    db.delete(doc)
    db.commit()
    db.refresh(claim)
    _audit(db, "DOCUMENT_DELETED", claim_id=cid, metadata={"document_id": str(did), "file_name": doc.file_name})
    logger.info("Deleted doc %s from claim %s", doc_id, claim_id)
    return _build_claim_response(db, cid)


@router.delete("/claims", status_code=204)
def delete_all_claims(db: Session = Depends(get_db)):
    # Delete all raw files from disk
    docs = db.query(Document).all()
    for doc in docs:
        try:
            p = Path(doc.minio_path).resolve()
            if str(p).startswith(str(RAW_STORAGE)):
                p.unlink(missing_ok=True)
        except OSError:
            logger.warning("Failed to delete file %s", doc.minio_path)

    db.query(Claim).delete()
    db.commit()
    logger.info("All claims deleted")


@router.delete("/claims/{claim_id}", status_code=204)
def delete_claim(claim_id: str, db: Session = Depends(get_db)):
    cid = _parse_uuid(claim_id)
    claim = db.query(Claim).filter(Claim.id == cid).first()
    if not claim:
        raise HTTPException(status_code=404, detail="Claim not found")

    # delete stored files from disk
    docs = db.query(Document).filter(Document.claim_id == cid).all()
    for doc in docs:
        try:
            p = Path(doc.minio_path).resolve()
            if str(p).startswith(str(RAW_STORAGE)):
                p.unlink(missing_ok=True)
        except OSError:
            logger.warning("Failed to delete file %s", doc.minio_path)

    doc_names = [d.file_name for d in docs]
    db.delete(claim)
    db.commit()
    _audit(db, "CLAIM_DELETED", claim_id=cid, metadata={"documents": doc_names})
    logger.info("Claim %s deleted", claim_id)

# ── Include router (standalone mode) ──
app.include_router(router)
